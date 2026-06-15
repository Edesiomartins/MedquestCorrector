import json
import re
from io import BytesIO

from fastapi import APIRouter, Depends, File, HTTPException, Query, UploadFile, status
from fastapi.responses import Response
from sqlalchemy import func
from sqlalchemy.orm import Session
from typing import List
from uuid import UUID
from docx import Document

from app.api.deps import get_current_user
from app.core.database import get_db
from app.models.exam import Exam, ExamQuestion
from app.models.grading import QuestionScore, StudentResult
from app.models.pipeline import UploadBatch
from app.models.student import Student
from app.models.user import Class
from app.schemas.exam import (
    ExamCreate,
    ExamQuestionCreate,
    ExamQuestionResponse,
    ExamResponse,
    ExamSummary,
)
from app.services.generator.answer_sheet import (
    QuestionSlot,
    StudentInfo,
    auto_fit_practical_sheet_options,
    generate_answer_sheets,
)

router = APIRouter(dependencies=[Depends(get_current_user)])
_DOCX_TEMPLATE_QUESTION_SLOTS = 10


@router.get("", response_model=List[ExamSummary])
def list_exams(
    practical: bool | None = Query(
        None,
        description="true=só provas práticas; false=só provas normais; omitido=lista todas",
    ),
    db: Session = Depends(get_db),
):
    counts = (
        db.query(ExamQuestion.exam_id, func.count(ExamQuestion.id))
        .group_by(ExamQuestion.exam_id)
        .all()
    )
    count_map = {eid: c for eid, c in counts}
    q = db.query(Exam)
    if practical is not None:
        q = q.filter(Exam.is_practical == practical)
    exams = q.order_by(Exam.name).all()
    return [
        ExamSummary(
            id=e.id,
            name=e.name,
            class_id=e.class_id,
            question_count=count_map.get(e.id, 0),
            is_practical=bool(e.is_practical),
        )
        for e in exams
    ]


@router.post("", response_model=ExamResponse, status_code=status.HTTP_201_CREATED)
def create_exam(exam_in: ExamCreate, db: Session = Depends(get_db)):
    new_exam = Exam(
        name=exam_in.name,
        class_id=exam_in.class_id,
        is_practical=exam_in.is_practical,
    )
    db.add(new_exam)
    db.commit()
    db.refresh(new_exam)
    return new_exam


@router.get("/templates/discursive-docx")
def download_discursive_docx_template():
    doc = Document()
    doc.add_heading("Template de Prova Discursiva", level=1)
    doc.add_paragraph(
        "Preencha as questões necessárias. Questões vazias serão ignoradas."
    )
    doc.add_paragraph("CURSO: ")
    doc.add_paragraph("DISCIPLINA: ")
    doc.add_paragraph("TURMA: ")
    doc.add_paragraph("VALOR GERAL: 1.0")
    doc.add_paragraph("")

    for idx in range(1, _DOCX_TEMPLATE_QUESTION_SLOTS + 1):
        doc.add_heading(f"QUESTÃO {idx}", level=2)
        doc.add_paragraph("Enunciado: ")
        doc.add_paragraph("Resposta esperada: ")
        if idx == 1:
            doc.add_paragraph(
                "Exemplo: descreva conceitos essenciais e critérios de pontuação de forma objetiva."
            )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="template_prova_discursiva.docx"'},
    )


@router.post("/import-discursive-docx")
async def import_discursive_docx(
    file: UploadFile = File(...),
    practical: bool = Query(False),
    db: Session = Depends(get_db),
):
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(
            status_code=400,
            detail={
                "ok": False,
                "message": "Não foi possível importar o DOCX.",
                "detail": "Envie um arquivo .docx válido.",
                "stage": "docx_import",
            },
        )
    raw = await file.read()
    if not raw:
        raise HTTPException(
            status_code=400,
            detail={
                "ok": False,
                "message": "Não foi possível importar o DOCX.",
                "detail": "Arquivo vazio.",
                "stage": "docx_import",
            },
        )

    try:
        parsed = _parse_discursive_docx(raw)
        title = parsed["metadata"].get("titulo") or _safe_title_from_filename(file.filename)
        turma_name = parsed["metadata"].get("turma")
        class_id = None
        if turma_name:
            turma = db.query(Class).filter(func.lower(Class.name) == turma_name.lower()).first()
            if turma:
                class_id = turma.id

        exam = Exam(name=title, class_id=class_id, is_practical=practical)
        db.add(exam)
        db.flush()

        warnings = list(parsed["warnings"])
        created = 0
        for q in parsed["questions"]:
            if not q["question_text"].strip():
                warnings.append(f"Questão {q['question_number']} ignorada por enunciado vazio.")
                continue
            db.add(
                ExamQuestion(
                    exam_id=exam.id,
                    question_number=q["question_number"],
                    question_text=q["question_text"].strip(),
                    expected_answer=(q["expected_answer"] or "Resposta esperada não informada.").strip(),
                    correction_criteria=(q.get("correction_criteria") or "").strip() or None,
                    max_score=float(q["max_score"]),
                )
            )
            created += 1

        if created == 0:
            db.rollback()
            raise ValueError("Nenhuma questão com enunciado foi encontrada.")

        db.commit()
        return {
            "ok": True,
            "exam_id": str(exam.id),
            "title": exam.name,
            "questions_created": created,
            "warnings": warnings,
        }
    except HTTPException:
        raise
    except Exception as exc:
        db.rollback()
        raise HTTPException(
            status_code=400,
            detail={
                "ok": False,
                "message": "Não foi possível importar o DOCX.",
                "detail": str(exc)[:500],
                "stage": "docx_import",
            },
        ) from exc


@router.get("/{exam_id}", response_model=ExamResponse)
def get_exam(exam_id: UUID, db: Session = Depends(get_db)):
    exam = db.query(Exam).filter(Exam.id == exam_id).first()
    if not exam:
        raise HTTPException(status_code=404, detail="Prova não encontrada.")
    return exam


@router.put("/{exam_id}", response_model=ExamResponse)
def update_exam(exam_id: UUID, exam_in: ExamCreate, db: Session = Depends(get_db)):
    exam = db.query(Exam).filter(Exam.id == exam_id).first()
    if not exam:
        raise HTTPException(status_code=404, detail="Prova não encontrada.")
    exam.name = exam_in.name
    exam.class_id = exam_in.class_id
    exam.is_practical = exam_in.is_practical
    db.commit()
    db.refresh(exam)
    return exam


@router.delete("/{exam_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_exam(exam_id: UUID, db: Session = Depends(get_db)):
    exam = db.query(Exam).filter(Exam.id == exam_id).first()
    if not exam:
        raise HTTPException(status_code=404, detail="Prova não encontrada.")
    try:
        question_ids = [
            qid
            for (qid,) in db.query(ExamQuestion.id).filter(ExamQuestion.exam_id == exam_id).all()
        ]
        batch_ids = [
            bid for (bid,) in db.query(UploadBatch.id).filter(UploadBatch.exam_id == exam_id).all()
        ]

        if question_ids:
            db.query(QuestionScore).filter(QuestionScore.question_id.in_(question_ids)).delete(
                synchronize_session=False
            )

        if batch_ids:
            result_ids = [
                rid
                for (rid,) in db.query(StudentResult.id).filter(
                    StudentResult.batch_id.in_(batch_ids)
                ).all()
            ]
            if result_ids:
                db.query(QuestionScore).filter(
                    QuestionScore.student_result_id.in_(result_ids)
                ).delete(synchronize_session=False)
                db.query(StudentResult).filter(StudentResult.id.in_(result_ids)).delete(
                    synchronize_session=False
                )

            db.query(UploadBatch).filter(UploadBatch.id.in_(batch_ids)).delete(
                synchronize_session=False
            )

        db.query(ExamQuestion).filter(ExamQuestion.exam_id == exam_id).delete(
            synchronize_session=False
        )
        db.delete(exam)
        db.commit()
    except Exception:
        db.rollback()
        raise HTTPException(
            status_code=409,
            detail="Não foi possível excluir a prova porque existem dados vinculados inconsistentes.",
        )


@router.delete("/{exam_id}/questions/{question_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_question(exam_id: UUID, question_id: UUID, db: Session = Depends(get_db)):
    q = db.query(ExamQuestion).filter(
        ExamQuestion.id == question_id, ExamQuestion.exam_id == exam_id
    ).first()
    if not q:
        raise HTTPException(status_code=404, detail="Questão não encontrada.")
    db.delete(q)
    db.commit()


@router.get("/{exam_id}/questions", response_model=List[ExamQuestionResponse])
def list_questions(exam_id: UUID, db: Session = Depends(get_db)):
    return (
        db.query(ExamQuestion)
        .filter(ExamQuestion.exam_id == exam_id)
        .order_by(ExamQuestion.question_number)
        .all()
    )


@router.post("/{exam_id}/questions", response_model=ExamQuestionResponse)
def add_question(exam_id: UUID, q_in: ExamQuestionCreate, db: Session = Depends(get_db)):
    exam = db.query(Exam).filter(Exam.id == exam_id).first()
    if not exam:
        raise HTTPException(status_code=404, detail="Prova não encontrada.")

    q = ExamQuestion(exam_id=exam_id, **q_in.model_dump())
    db.add(q)
    db.commit()
    db.refresh(q)
    return q


@router.put("/{exam_id}/questions/{question_id}", response_model=ExamQuestionResponse)
def update_question(
    exam_id: UUID,
    question_id: UUID,
    q_in: ExamQuestionCreate,
    db: Session = Depends(get_db),
):
    q = db.query(ExamQuestion).filter(
        ExamQuestion.id == question_id, ExamQuestion.exam_id == exam_id
    ).first()
    if not q:
        raise HTTPException(status_code=404, detail="Questão não encontrada.")

    q.question_number = q_in.question_number
    q.question_text = q_in.question_text
    q.expected_answer = q_in.expected_answer
    q.correction_criteria = q_in.correction_criteria
    q.max_score = q_in.max_score
    db.commit()
    db.refresh(q)
    return q


@router.get("/{exam_id}/answer-sheets")
def download_answer_sheets(exam_id: UUID, db: Session = Depends(get_db)):
    return _build_answer_sheets_response(exam_id=exam_id, db=db)


@router.post("/{exam_id}/answer-sheets")
def download_answer_sheets_with_logo(
    exam_id: UUID,
    logo: UploadFile | None = File(default=None),
    db: Session = Depends(get_db),
):
    logo_bytes = _read_logo_bytes(logo)
    return _build_answer_sheets_response(exam_id=exam_id, db=db, logo_bytes=logo_bytes)


@router.post("/{exam_id}/answer-sheets/practical")
def download_practical_answer_sheets_with_logo(
    exam_id: UUID,
    logo: UploadFile | None = File(default=None),
    db: Session = Depends(get_db),
):
    logo_bytes = _read_logo_bytes(logo)
    return _build_answer_sheets_response(
        exam_id=exam_id,
        db=db,
        logo_bytes=logo_bytes,
        auto_fit_practical=True,
    )


def _build_answer_sheets_response(
    exam_id: UUID,
    db: Session,
    logo_bytes: bytes | None = None,
    sheet_options: dict | None = None,
    auto_fit_practical: bool = False,
):
    """Gera e baixa folhas-resposta PDF para todos os alunos da turma vinculada."""
    exam = db.query(Exam).filter(Exam.id == exam_id).first()
    if not exam:
        raise HTTPException(status_code=404, detail="Prova não encontrada.")

    if not exam.class_id:
        raise HTTPException(status_code=400, detail="Prova não vinculada a uma turma.")

    turma = db.query(Class).filter(Class.id == exam.class_id).first()
    turma_name = turma.name if turma else "—"

    students = (
        db.query(Student)
        .filter(Student.class_id == exam.class_id)
        .order_by(Student.registration_number)
        .all()
    )
    if not students:
        raise HTTPException(status_code=404, detail="Nenhum aluno na turma.")

    questions = (
        db.query(ExamQuestion)
        .filter(ExamQuestion.exam_id == exam_id)
        .order_by(ExamQuestion.question_number)
        .all()
    )
    if not questions:
        raise HTTPException(status_code=400, detail="Prova sem questões cadastradas.")

    question_slots = [
        QuestionSlot(number=q.question_number, text=q.question_text, max_score=q.max_score)
        for q in questions
    ]
    effective_sheet_options = sheet_options
    if auto_fit_practical:
        effective_sheet_options = auto_fit_practical_sheet_options(
            question_slots,
            has_logo=logo_bytes is not None,
        )

    try:
        pdf_bytes, manifest = generate_answer_sheets(
            exam_id=exam_id,
            exam_name=exam.name,
            questions=question_slots,
            students=[
                (
                    s.id,
                    StudentInfo(
                        name=s.name,
                        registration_number=s.registration_number,
                        curso=s.curso or "—",
                        turma=turma_name,
                    ),
                )
                for s in students
            ],
            logo_bytes=logo_bytes,
            **(effective_sheet_options or {}),
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    exam.layout_manifest_json = json.dumps(manifest, ensure_ascii=False)
    db.commit()

    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="folhas_{exam.name}.pdf"'},
    )


def _parse_discursive_docx(raw: bytes) -> dict:
    doc = Document(BytesIO(raw))
    lines = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    metadata: dict[str, str] = {}
    questions: dict[int, dict] = {}
    warnings: list[str] = []
    current_q: int | None = None
    current_field: str | None = None

    default_score = 1.0
    for line in lines:
        q_match = re.match(r"^(?:quest[aã]o|q)\s*(\d+)\b", line, flags=re.IGNORECASE)
        if q_match:
            qnum = int(q_match.group(1))
            if qnum >= 1:
                current_q = qnum
                current_field = None
                questions.setdefault(
                    qnum,
                    {
                        "question_number": qnum,
                        "question_text": "",
                        "expected_answer": "",
                        "correction_criteria": "",
                        "max_score": default_score,
                    },
                )
            continue

        key, value = _split_docx_field(line)
        if key:
            normalized = _normalize_docx_key(key)
            if normalized == "valor_padrao":
                default_score = _parse_score(value, default_score)
                for q in questions.values():
                    if not q.get("max_score"):
                        q["max_score"] = default_score
                metadata[normalized] = str(default_score)
                continue
            if normalized in {"titulo", "disciplina", "curso", "turma", "instrucoes"} and current_q is None:
                metadata[normalized] = value.strip()
                continue
            if current_q is not None and normalized in {
                "question_text",
                "expected_answer",
                "correction_criteria",
                "max_score",
            }:
                current_field = normalized
                if normalized == "max_score":
                    questions[current_q][normalized] = _parse_score(value, default_score)
                else:
                    questions[current_q][normalized] = _append_text(questions[current_q].get(normalized, ""), value)
                continue

        if current_q is not None and current_field:
            if current_field == "max_score":
                questions[current_q][current_field] = _parse_score(line, default_score)
            else:
                questions[current_q][current_field] = _append_text(questions[current_q].get(current_field, ""), line)

    parsed_questions = []
    for qnum in sorted(questions):
        q = questions[qnum]
        if not q["question_text"].strip():
            warnings.append(f"Questão {qnum} ignorada: enunciado vazio.")
            continue
        if not q["expected_answer"].strip():
            warnings.append(f"Questão {qnum}: resposta esperada não informada.")
        parsed_questions.append(q)

    return {"metadata": metadata, "questions": parsed_questions, "warnings": warnings}


def _split_docx_field(line: str) -> tuple[str | None, str]:
    if ":" not in line:
        return None, line
    key, value = line.split(":", 1)
    return key.strip(), value.strip()


def _normalize_docx_key(key: str) -> str:
    raw = key.strip().lower()
    raw = raw.replace("í", "i").replace("ç", "c").replace("ã", "a").replace("õ", "o")
    if raw in {"titulo da prova", "titulo", "título da prova", "título"}:
        return "titulo"
    if raw == "disciplina":
        return "disciplina"
    if raw == "curso":
        return "curso"
    if raw == "turma":
        return "turma"
    if raw.startswith("instrucoes"):
        return "instrucoes"
    if raw.startswith("valor padrao") or raw.startswith("valor geral"):
        return "valor_padrao"
    if raw in {"enunciado", "pergunta"}:
        return "question_text"
    if raw in {"resposta esperada", "gabarito", "padrao de resposta", "padrão de resposta"}:
        return "expected_answer"
    if raw in {"criterios", "criterios de correcao", "critérios", "critérios de correção", "rubrica"}:
        return "correction_criteria"
    if raw in {"valor", "pontuacao", "pontuação"}:
        return "max_score"
    return raw


def _parse_score(value: str, default: float) -> float:
    try:
        return float(str(value).replace(",", ".").strip())
    except (TypeError, ValueError):
        return default


def _append_text(existing: str, value: str) -> str:
    value = (value or "").strip()
    if not value:
        return existing or ""
    if not existing:
        return value
    return f"{existing}\n{value}"


def _safe_title_from_filename(filename: str) -> str:
    name = (filename or "Prova discursiva").rsplit(".", 1)[0]
    return name.replace("_", " ").strip() or "Prova discursiva"


def _read_logo_bytes(logo: UploadFile | None) -> bytes | None:
    if not logo or not logo.filename:
        return None

    allowed_types = {"image/png", "image/jpeg", "image/jpg", "image/webp"}
    if not logo.content_type or logo.content_type.lower() not in allowed_types:
        raise HTTPException(
            status_code=400,
            detail="A logo deve estar em PNG, JPG ou WEBP.",
        )

    logo_bytes = logo.file.read()
    if not logo_bytes:
        raise HTTPException(status_code=400, detail="A logo enviada está vazia.")

    max_size = 5 * 1024 * 1024
    if len(logo_bytes) > max_size:
        raise HTTPException(status_code=400, detail="A logo deve ter no máximo 5 MB.")

    return logo_bytes
