from io import BytesIO

from docx import Document
from openpyxl import load_workbook

from app.api.v1.exams import _parse_discursive_docx
from app.services.export.spreadsheet import export_results_xlsx, humanize_review_reason


def test_parse_discursive_docx_with_three_questions():
    doc = Document()
    doc.add_paragraph("TÍTULO DA PROVA: Fisiologia")
    doc.add_paragraph("TURMA: MED")
    doc.add_paragraph("VALOR PADRÃO POR QUESTÃO: 1")
    for idx in range(1, 4):
        doc.add_paragraph(f"Q{idx}")
        doc.add_paragraph(f"Enunciado: Enunciado {idx}")
        doc.add_paragraph(f"Gabarito: Resposta {idx}")
        doc.add_paragraph(f"Rubrica: Critério {idx}")
        doc.add_paragraph("Pontuação: 1")
    buf = BytesIO()
    doc.save(buf)

    parsed = _parse_discursive_docx(buf.getvalue())

    assert parsed["metadata"]["titulo"] == "Fisiologia"
    assert len(parsed["questions"]) == 3
    assert parsed["questions"][1]["expected_answer"] == "Resposta 2"
    assert parsed["questions"][2]["correction_criteria"] == "Critério 3"


def test_parse_discursive_docx_without_hard_limit_of_questions():
    doc = Document()
    doc.add_paragraph("TÍTULO DA PROVA: Farmacologia")
    doc.add_paragraph("TURMA: MED")
    doc.add_paragraph("VALOR PADRÃO POR QUESTÃO: 1")
    for idx in range(1, 13):
        doc.add_paragraph(f"Q{idx}")
        doc.add_paragraph(f"Enunciado: Enunciado {idx}")
        doc.add_paragraph(f"Gabarito: Resposta {idx}")
    buf = BytesIO()
    doc.save(buf)

    parsed = _parse_discursive_docx(buf.getvalue())

    assert len(parsed["questions"]) == 12
    assert parsed["questions"][-1]["question_number"] == 12


def test_humanize_review_reason_hides_technical_json_from_main_reason():
    reason = humanize_review_reason(["JSON recuperado, mas schema incompleto"])
    assert reason == "Falha no formato da resposta da IA. Conferir nota sugerida."


def test_export_results_xlsx_summary_student_and_review_sheets():
    data = export_results_xlsx(
        "Prova",
        [{"number": 1, "max_score": 1}, {"number": 2, "max_score": 1}, {"number": 3, "max_score": 1}],
        [
            {
                "registration_number": "001",
                "student_name": "Aluno 1",
                "curso": "Med",
                "turma": "T1",
                "scores": {1: 0.5, 2: 0.5, 3: 0.5},
                "needs_review": True,
                "warnings": ["JSON recuperado, mas schema incompleto"],
                "question_details": [
                    {
                        "question_number": 1,
                        "score": 0.5,
                        "verdict": "parcial",
                        "comment": "Comentário",
                        "transcription": "Texto",
                        "needs_review": True,
                        "review_reason": "Expecting property name enclosed in double quotes",
                        "technical_detail": "Expecting property name enclosed in double quotes",
                    }
                ],
            }
        ],
    )
    wb = load_workbook(BytesIO(data))
    assert wb.sheetnames[0] == "Resultado Final"
    assert wb.sheetnames[-1] == "Revisões Necessárias"
    assert len(wb.sheetnames) == 3
    ws = wb["Resultado Final"]
    assert ws["E3"].value == 0.5
    assert "JSON recuperado" not in str(ws["G3"].value)
    alum = wb[wb.sheetnames[1]]
    assert alum["A5"].value == "Matrícula"
    assert alum["B5"].value == "001"
    assert alum["A17"].value == "Questão"
    assert alum["A18"].value == 1
    assert alum["C18"].value == "parcial"

